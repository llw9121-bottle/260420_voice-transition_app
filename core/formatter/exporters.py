"""
文档导出模块

支持多种格式的文档导出：Word (.docx)、Markdown (.md)、JSON (.json)。
"""

import json
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from loguru import logger

from core.formatter.base import FormattedDocument, BehaviorMatch


class BaseExporter(ABC):
    """导出器基类"""
    
    @property
    @abstractmethod
    def file_extension(self) -> str:
        """文件扩展名"""
        pass
    
    @abstractmethod
    def export(self, document: FormattedDocument, output_path: Path) -> Path:
        """
        导出文档
        
        Args:
            document: 格式化后的文档
            output_path: 输出路径
            
        Returns:
            导出的文件路径
        """
        pass
    
    def _ensure_extension(self, path: Path) -> Path:
        """确保路径有正确的扩展名"""
        if path.suffix.lower() != self.file_extension.lower():
            return path.with_suffix(self.file_extension)
        return path


class JSONExporter(BaseExporter):
    """JSON导出器"""
    
    @property
    def file_extension(self) -> str:
        return ".json"
    
    def export(self, document: FormattedDocument, output_path: Path) -> Path:
        output_path = self._ensure_extension(output_path)
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 转换为字典
        data = document.to_dict()
        
        # 写入JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"JSON文档已导出: {output_path}")
        return output_path


class MarkdownExporter(BaseExporter):
    """Markdown导出器"""
    
    @property
    def file_extension(self) -> str:
        return ".md"
    
    def export(self, document: FormattedDocument, output_path: Path) -> Path:
        output_path = self._ensure_extension(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 构建Markdown内容
        lines = []
        
        # 标题
        title = document.title or "转录文档"
        lines.append(f"# {title}")
        lines.append("")
        
        # 元信息
        lines.append(f"**创建时间**: {document.created_at.strftime('%Y-%m-%d %H:%M')}")
        if document.session_id:
            lines.append(f"**会话ID**: {document.session_id}")
        lines.append(f"**格式化风格**: {document.style.value}")
        lines.append("")
        
        # 正文
        lines.append("## 正文")
        lines.append("")
        
        if document.formatted_text:
            # 处理带标记的文本，将标记转换为高亮格式
            formatted = self._format_behavior_marks(document.formatted_text)
            lines.append(formatted)
        else:
            lines.append(document.raw_text)
        
        lines.append("")
        
        # 行为匹配结果
        if document.behavior_matches:
            lines.append("## 行为匹配结果")
            lines.append("")
            
            for i, match in enumerate(document.behavior_matches, 1):
                lines.append(f"### {i}. {match.behavior_name}")
                lines.append(f"- **置信度**: {match.confidence:.1%}")
                lines.append(f"- **原文引用**: \"{match.original_text}\"")
                lines.append("")
        
        # 统计信息
        lines.append("---")
        lines.append("")
        lines.append("## 统计信息")
        lines.append("")
        lines.append(f"- **字数**: {document.word_count}")
        lines.append(f"- **时长**: {document.duration_seconds:.1f}秒")
        if document.speaker_count > 0:
            lines.append(f"- **说话人数量**: {document.speaker_count}")
        
        # 写入文件
        content = '\n'.join(lines)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        logger.info(f"Markdown文档已导出: {output_path}")
        return output_path
    
    def _format_behavior_marks(self, text: str) -> str:
        """将行为标记转换为Markdown高亮格式"""
        import re
        
        # 匹配【行为名(置信度)】模式
        pattern = r'【([^】]+)\((\d+%)\)】'
        
        def replace_mark(match):
            behavior = match.group(1)
            confidence = match.group(2)
            return f" **`[{behavior}|{confidence}]`** "
        
        return re.sub(pattern, replace_mark, text)


class WordExporter(BaseExporter):
    """Word文档导出器"""
    
    @property
    def file_extension(self) -> str:
        return ".docx"
    
    def export(self, document: FormattedDocument, output_path: Path) -> Path:
        output_path = self._ensure_extension(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise ImportError(
                "导出Word文档需要python-docx库。"
                "请运行: pip install python-docx"
            )
        
        # 创建文档
        doc = Document()
        
        # 设置默认字体
        # python-docx 需要单独设置东亚字体才能让中文正确显示
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        # 设置东亚字体，这是中文正确显示的关键
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
        
        # 标题
        title = document.title or "转录文档"
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 元信息表格
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        meta_data = [
            ("创建时间", document.created_at.strftime('%Y-%m-%d %H:%M')),
            ("格式化风格", document.style.value),
            ("字数", str(document.word_count)),
        ]
        
        for i, (key, value) in enumerate(meta_data):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
        
        doc.add_paragraph()  # 空行
        
        # 正文标题
        doc.add_heading('正文', level=2)
        
        # 正文内容（处理行为标记）
        if document.formatted_text:
            self._add_formatted_content(doc, document.formatted_text)
        else:
            doc.add_paragraph(document.raw_text)
        
        # 行为匹配结果
        if document.behavior_matches:
            doc.add_page_break()
            doc.add_heading('行为匹配结果', level=2)
            
            for i, match in enumerate(document.behavior_matches, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(match.behavior_name).bold = True
                
                info = doc.add_paragraph(style='List Bullet')
                info.add_run(f"置信度: {match.confidence:.1%}")
                
                quote = doc.add_paragraph(style='List Bullet')
                quote.add_run('原文引用: ').italic = True
                quote.add_run(f'"{match.original_text}"')
        
        # 统计信息
        doc.add_paragraph()
        doc.add_heading('统计信息', level=2)
        
        stats = doc.add_paragraph()
        stats.add_run(f"字数: {document.word_count}\n")
        stats.add_run(f"时长: {document.duration_seconds:.1f}秒\n")
        if document.speaker_count > 0:
            stats.add_run(f"说话人数量: {document.speaker_count}")
        
        # 保存文档
        doc.save(str(output_path))
        logger.info(f"Word文档已导出: {output_path}")
        
        return output_path
    
    def _add_formatted_content(self, doc, text: str):
        """添加格式化的内容，处理行为标记"""
        import re

        # 匹配【行为名(置信度)】整个标记
        # 只匹配标记本身，内容已经在格式化文本中，不要重复消费
        pattern = r'【([^】]+)\((\d+%)\)】'
        last_pos = 0

        current_para = doc.add_paragraph()

        for match in re.finditer(pattern, text):
            # 匹配前的普通文本
            if match.start() > last_pos:
                prev_text = text[last_pos:match.start()]
                if prev_text:
                    if '\n' in prev_text:
                        # 处理换行
                        lines = prev_text.split('\n')
                        for j, line in enumerate(lines):
                            if line:
                                current_para.add_run(line)
                            if j < len(lines) - 1:
                                current_para = doc.add_paragraph()
                    else:
                        current_para.add_run(prev_text)

            # 行为标记 - 将【行为名(置信度)】转换为加粗格式
            behavior = match.group(1)
            confidence = match.group(2)

            # 添加行为标记（加粗）
            run = current_para.add_run(f"[{behavior}|{confidence}] ")
            run.bold = True

            last_pos = match.end()

        # 添加剩余文本
        if last_pos < len(text):
            remaining = text[last_pos:]
            if remaining:
                if '\n' in remaining:
                    lines = remaining.split('\n')
                    for j, line in enumerate(lines):
                        if line:
                            current_para.add_run(line)
                        if j < len(lines) - 1:
                            current_para = doc.add_paragraph()
                else:
                    current_para.add_run(remaining)
